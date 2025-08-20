# modules/loader.py
from pathlib import Path
from typing import List, Dict
import sys

def _extract_pdf_text(path: str) -> str:
    """Try several PDF text extractors; if none produce text, fall back to OCR."""
    # 1) pdfplumber (pdfminer)
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            txt = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
        if txt:
            print("[loader/pdf] pdfplumber ok")
            return txt
        print("[loader/pdf] pdfplumber empty, trying pypdf…", file=sys.stderr)
    except Exception as e:
        print(f"[loader/pdf] pdfplumber failed: {e}", file=sys.stderr)

    # 2) PyPDF (pure text)
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        txt = "\n".join((p.extract_text() or "") for p in reader.pages).strip()
        if txt:
            print("[loader/pdf] pypdf ok")
            return txt
        print("[loader/pdf] pypdf empty, trying PyMuPDF…", file=sys.stderr)
    except Exception as e:
        print(f"[loader/pdf] pypdf failed: {e}", file=sys.stderr)

    # 3) PyMuPDF (fitz) text extraction
    import fitz
    doc = fitz.open(path)
    try:
        txt = "\n".join(page.get_text() or "" for page in doc).strip()
        if txt:
            print("[loader/pdf] pymupdf ok")
            return txt
        print("[loader/pdf] pymupdf empty, falling back to OCR…", file=sys.stderr)
    finally:
        doc.close()

    # 4) OCR fallback: render pages -> Tesseract
  # 4) OCR fallback: render pages -> Tesseract
    try:
        import pytesseract
        import pypdfium2 as pdfium
        from PIL import Image

        pdf = pdfium.PdfDocument(path)
        try:
            n = len(pdf)
            ocr_texts = []
            for i in range(n):
                page = pdf[i]
                # render at 300 DPI for better accuracy (was 200)
                bitmap = page.render(scale=300/72)     # returns PdfBitmap
                try:
                    pil_image = bitmap.to_pil()       # Pillow Image
                    try:
                        if pil_image.mode != "L":
                            pil_image = pil_image.convert("L")
                        txt = pytesseract.image_to_string(pil_image)
                        ocr_texts.append(txt.strip())
                    finally:
                        pil_image.close()
                finally:
                    # Explicitly drop bitmap before next page
                    del bitmap
            txt = "\n".join(ocr_texts).strip()
            if txt:
                print(f"[loader/pdf] OCR ok ({n} page(s))")
                return txt
            print("[loader/pdf] OCR produced empty text", file=sys.stderr)
        finally:
            # Always close the document
            pdf.close()
    except Exception as e:
        print(f"[loader/pdf] OCR path failed: {e}", file=sys.stderr)


    return ""  # give up

def load_files(folder: str, content_types: List[str]) -> List[Dict]:
    """
    Return a list of {"text": str, "metadata": {...}} for supported files in folder.
    - content_types: list of extensions without dot (e.g. ["pdf","txt","md","rtf","csv","docx"])
    - Case-insensitive matching
    - Logs what it sees so debugging is easy
    """
    p = Path(folder)
    if not p.exists():
        print(f"[loader] Folder not found: {folder}", file=sys.stderr)
        return []

    allowed = {ext.lower().lstrip(".") for ext in (content_types or [])}
    if not allowed:
        allowed = {"txt","md","pdf","rtf","csv","docx"}

    all_files = [f for f in p.iterdir() if f.is_file()]
    print(f"[loader] Found {len(all_files)} files in {folder}: {[f.name for f in all_files]}")
    cand = [f for f in all_files if f.suffix.lower().lstrip(".") in allowed]
    print(f"[loader] Candidates by type {sorted(allowed)}: {[f.name for f in cand]}")

    out = []
    for f in cand:
        ext = f.suffix.lower().lstrip(".")
        text = ""
        try:
            if ext in {"txt","md","rtf","csv"}:
                text = f.read_text(encoding="utf-8", errors="ignore")
            elif ext == "docx":
                try:
                    import docx
                    doc = docx.Document(str(f))
                    text = "\n".join(p.text for p in doc.paragraphs)
                except Exception as e:
                    print(f"[loader] python-docx failed on {f.name}: {e}", file=sys.stderr)
                    continue
            elif ext == "pdf":
                text = _extract_pdf_text(str(f))
            else:
                print(f"[loader] Skipping unsupported extension: {f.name}", file=sys.stderr)
                continue

            text = (text or "").strip()
            if not text:
                print(f"[loader] Empty text after parsing: {f.name}", file=sys.stderr)
                continue

            out.append({
                "text": text,
                "metadata": {"filename": f.name, "path": str(f), "ext": ext}
            })
        except Exception as e:
            print(f"[loader] Failed to load {f.name}: {e}", file=sys.stderr)
            continue

    print(f"[loader] Loaded {len(out)} document(s).")
    return out

